from pathlib import Path
from typing import Optional
import io
from collections import deque
from datetime import datetime

from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware

from strategies import load_strategies
from reports.report_service import generate_pdf

import pytesseract
from PIL import Image, UnidentifiedImageError

try:
    import fitz  # PyMuPDF for PDFs
except Exception:
    fitz = None
try:
    from docx import Document as DocxDocument  # python-docx
except Exception:
    DocxDocument = None

#------------memory store+helper ---------
SCAN_MEM = deque(maxlen=50)

def _push_mem_log(user: str, strategy: str, status: str) -> None:
    SCAN_MEM.appendleft({
        "ts": datetime.now().astimezone().isoformat(),
        "user": user,
        "strategy": strategy,
        "status": status,  # 'success' | 'error'
    })

# -------- extraction helpers --------
def _ocr_image_bytes(data: bytes) -> str:
    try:
        with Image.open(io.BytesIO(data)) as img:
            return pytesseract.image_to_string(img)
    except (UnidentifiedImageError, OSError):
        return ""


def _extract_pdf_bytes(data: bytes, previews_dir: Path, stem: str) -> tuple[str, Optional[Path]]:
    if not fitz:
        return "", None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        text = "".join(page.get_text() or "" for page in doc)
        preview_path = None
        if len(doc) > 0:
            pix = doc[0].get_pixmap()
            previews_dir.mkdir(parents=True, exist_ok=True)
            preview_path = previews_dir / f"{stem}_page1.png"
            pix.save(str(preview_path))
        doc.close()
        return text, preview_path
    except Exception:
        return "", None


def _extract_docx_bytes(data: bytes) -> str:
    if not DocxDocument:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_and_preview_bytes(filename: str, data: bytes, previews_dir: Path) -> tuple[str, Optional[Path]]:
    ext = Path(filename).suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        text = _ocr_image_bytes(data)
        previews_dir.mkdir(parents=True, exist_ok=True)
        preview_path = previews_dir / filename
        with open(preview_path, "wb") as f:
            f.write(data)
        return text, preview_path
    if ext == ".pdf":
        return _extract_pdf_bytes(data, previews_dir, Path(filename).stem)
    if ext == ".docx":
        return _extract_docx_bytes(data), None
    if ext in {".txt", ".log", ".reg", ".csv", ".ini", ".json", ".xml", ".htm", ".html"}:
        return data.decode("utf-8", errors="ignore"), None
    return "", None


# -------- FastAPI app --------
app = FastAPI(title="AutoAudit Evidence Scanner")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "templates"
OUT_DIR = ROOT / "reports_out"
PREVIEWS = ROOT / "previews"
INDEX_HTML = ROOT / "aa_ui" / "ui.html"  # <-- serve ui.html


@app.get("/", response_class=HTMLResponse)
def index():
    if INDEX_HTML.exists():
        return INDEX_HTML.read_text(encoding="utf-8", errors="ignore")
    return "<h3>AutoAudit Evidence Scanner</h3>"


@app.get("/strategies")
def api_strategies():
    out = []
    for s in load_strategies():
        try:
            desc = s.description()
        except Exception:
            desc = ""
        out.append({"name": s.name, "description": desc})
    return out

@app.post("/api/scan-mem-log")
def api_post_scan_mem_log(payload: dict):
    _push_mem_log(
        user=str(payload.get("user") or payload.get("user_id") or "user"),
        strategy=str(payload.get("strategy") or payload.get("strategy_name") or ""),
        status=str(payload.get("status") or "success"),
    )
    return {"ok": True}

@app.get("/api/scan-mem-log")
def api_get_scan_mem_log():
    return JSONResponse(list(SCAN_MEM))

@app.get("/scan-mem", response_class=HTMLResponse)
def scan_mem_page():
    rows = "".join(
        f"<tr><td>{r.get('ts','')}</td>"
        f"<td>{r.get('user','')}</td>"
        f"<td>{r.get('strategy','')}</td>"
        f"<td>{r.get('status','')}</td></tr>"
        for r in SCAN_MEM
    ) or "<tr><td colspan='4'>No runs yet</td></tr>"
    return HTMLResponse(f"""<!doctype html><meta charset="utf-8">
<h1>Recent Scans — in memory</h1>
<p><a href="/">← Home</a></p>
<table border="1" cellpadding="6">
  <tr><th>Time</th><th>User</th><th>Strategy</th><th>Status</th></tr>
  {rows}
</table>""")



@app.post("/scan")
async def scan(
    evidence: UploadFile = File(...),
    strategy_name: str = Form(...),
    user_id: str = Form("user"),
):
    # find strategy
    strategy = next((s for s in load_strategies() if s.name == strategy_name), None)
    if not strategy:
        raise HTTPException(status_code=400, detail=f"Unknown strategy: {strategy_name}")

    # read upload
    content = await evidence.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty upload")

    # extract text + preview
    text, preview_path = extract_text_and_preview_bytes(evidence.filename, content, PREVIEWS)
    if not text.strip():
        return JSONResponse({"ok": True, "findings": [], "reports": [], "note": "No readable text found in evidence."})

    # run rules
    if hasattr(strategy, "emit_hits"):
        findings = strategy.emit_hits(text, source_file=evidence.filename) or []
    else:
        hits = strategy.match(text) or []
        findings = [{
            "test_id": "",
            "sub_strategy": "",
            "detected_level": "",
            "pass_fail": "",
            "priority": "",
            "recommendation": "",
            "evidence": hits,
        }] if hits else []

    # generate PDFs
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = []
    for r in findings:
        data = {
            "UniqueID": f"{user_id}-{strategy.name}-{Path(evidence.filename).stem}",
            "UserID": user_id,
            "Evidence": evidence.filename,
            "Evidence Preview": str(preview_path) if preview_path else "",
            "Strategy": strategy.name,
            "TestID": r.get("test_id", ""),
            "Sub-Strategy": r.get("sub_strategy", ""),
            "ML Level": r.get("detected_level", ""),
            "Pass/Fail": r.get("pass_fail", ""),
            "Priority": r.get("priority", ""),
            "Recommendation": r.get("recommendation", ""),
            "Evidence Extract": "; ".join(r.get("evidence", [])),
            "Description": r.get("description", ""),
            "Confidence": r.get("confidence", ""),
        }
        pdf_path = generate_pdf(
            data,
            template_path=str(TEMPLATES / "report_template.docx"),
            output_dir=str(OUT_DIR),
            base_dir=str(ROOT),
        )
        generated.append(Path(pdf_path).name)

    return {"ok": True, "findings": findings, "reports": generated}


@app.get("/reports/{filename}")
def download_report(filename: str):
    path = OUT_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    return FileResponse(str(path), media_type="application/pdf", filename=filename)
